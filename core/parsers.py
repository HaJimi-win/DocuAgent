"""
文档解析模块：多格式文档解析 + MD5增量去重
支持格式：TXT、CSV、PDF、Word(.docx)、Excel(.xlsx/.xls)
"""
import os
import json
import hashlib
from fastapi import HTTPException
import pandas as pd

from config import (
    WORKSPACE, UPLOAD_RECORD, text_splitter,
    SUPPORTED_EXTENSIONS, logger,
)

# 多格式解析库（可选导入）
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================
# MD5 增量去重
# ============================================================
def calculate_md5(file_path: str) -> str:
    """计算文件MD5哈希（流式读取，大文件不占内存）"""
    md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            md5.update(chunk)
    return md5.hexdigest()


def load_upload_records() -> dict:
    """加载已上传文件记录"""
    if os.path.exists(UPLOAD_RECORD):
        try:
            with open(UPLOAD_RECORD, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.warning("读取上传记录失败，将重建: %s", e)
    return {}


def save_upload_record(records: dict):
    """保存上传记录"""
    try:
        with open(UPLOAD_RECORD, "w", encoding="utf-8") as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("保存上传记录失败: %s", e)


# ============================================================
# 各格式解析器
# ============================================================
def parse_txt(file_path: str) -> list:
    """解析TXT：流式按行读取，累积到块大小后切片"""
    chunks = []
    buffer = ""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                buffer += line
                if len(buffer) >= 2000:
                    chunks.extend(text_splitter.split_text(buffer))
                    buffer = ""
        if buffer.strip():
            chunks.extend(text_splitter.split_text(buffer))
        logger.info("TXT解析完成: %s, 切片数: %d", os.path.basename(file_path), len(chunks))
        return chunks
    except UnicodeDecodeError as e:
        logger.error("TXT编码错误: %s", e)
        raise HTTPException(status_code=400, detail=f"文件编码错误，请使用UTF-8编码: {e}")


def parse_csv(file_path: str) -> list:
    """解析CSV：用pandas读取，转成文本描述（列名+前N行+统计摘要）"""
    try:
        df = pd.read_csv(file_path)
        lines = [f"CSV文件列名: {list(df.columns)}", f"数据行数: {len(df)}, 列数: {len(df.columns)}"]
        lines.append(f"前5行数据:\n{df.head(5).to_string()}")
        try:
            lines.append(f"数值列统计摘要:\n{df.describe().to_string()}")
        except Exception:
            pass
        text = "\n".join(lines)
        chunks = text_splitter.split_text(text)
        logger.info("CSV解析完成: %s, 切片数: %d", os.path.basename(file_path), len(chunks))
        return chunks
    except Exception as e:
        logger.error("CSV解析失败: %s", e)
        raise HTTPException(status_code=400, detail=f"CSV解析失败: {e}")


def parse_pdf(file_path: str) -> list:
    """解析PDF：逐页提取文本（天然流式）"""
    if not HAS_PYPDF:
        raise HTTPException(status_code=400, detail="未安装pypdf，无法解析PDF，请执行: pip install pypdf")
    try:
        reader = PdfReader(file_path)
        chunks = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                page_chunks = text_splitter.split_text(text)
                chunks.extend(page_chunks)
            if (i + 1) % 20 == 0:
                logger.info("PDF解析进度: 第%d/%d页", i + 1, len(reader.pages))
        logger.info("PDF解析完成: %s, 总页数: %d, 切片数: %d",
                    os.path.basename(file_path), len(reader.pages), len(chunks))
        return chunks
    except Exception as e:
        logger.error("PDF解析失败: %s", e)
        raise HTTPException(status_code=400, detail=f"PDF解析失败: {e}")


def parse_docx(file_path: str) -> list:
    """解析Word(.docx)：逐段落提取文本 + 表格内容"""
    if not HAS_DOCX:
        raise HTTPException(status_code=400, detail="未安装python-docx，无法解析Word，请执行: pip install python-docx")
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)
        text = "\n".join(full_text)
        chunks = text_splitter.split_text(text)
        logger.info("Word解析完成: %s, 切片数: %d", os.path.basename(file_path), len(chunks))
        return chunks
    except Exception as e:
        logger.error("Word解析失败: %s", e)
        raise HTTPException(status_code=400, detail=f"Word解析失败: {e}")


def parse_excel(file_path: str) -> list:
    """解析Excel(.xlsx/.xls)：每个sheet转文本描述"""
    try:
        xls = pd.ExcelFile(file_path)
        all_text = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            all_text.append(f"=== Sheet: {sheet_name} ===")
            all_text.append(f"列名: {list(df.columns)}, 行数: {len(df)}")
            all_text.append(f"前5行数据:\n{df.head(5).to_string()}")
            try:
                all_text.append(f"统计摘要:\n{df.describe().to_string()}")
            except Exception:
                pass
        text = "\n".join(all_text)
        chunks = text_splitter.split_text(text)
        logger.info("Excel解析完成: %s, sheet数: %d, 切片数: %d",
                    os.path.basename(file_path), len(xls.sheet_names), len(chunks))
        return chunks
    except Exception as e:
        logger.error("Excel解析失败: %s", e)
        raise HTTPException(status_code=400, detail=f"Excel解析失败: {e}")


# 解析器分发字典
PARSERS = {
    ".txt": parse_txt,
    ".csv": parse_csv,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
}


def parse_document(file_path: str) -> list:
    """根据文件扩展名分发到对应解析器，返回文本切片列表"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in PARSERS:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}，支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
    return PARSERS[ext](file_path)
